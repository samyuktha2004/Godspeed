from __future__ import annotations

import logging

from src.file_agent.parsers import Block, register

logger = logging.getLogger(__name__)


@register("docx")
def parse_docx(path: str) -> list[Block]:
    blocks: list[Block] = []
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed — cannot parse DOCX")
        return blocks

    doc = Document(path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append({"type": "text", "content": text, "style": para.style.name})

    for tbl_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            blocks.append({"type": "table", "content": rows, "table_index": tbl_idx})

    return blocks
